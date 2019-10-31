# -*- coding：utf-8 -*-
# &Author  AnFany


# 将.docx版的文档转变为pdf格式


from docx import Document
from docx.shared import Inches, RGBColor
from docx.oxml.ns import qn
from docx.shared import Pt
import docx.enum.text as s
import baby_grow_report as g_report
from win32com import client  # docx转化为pdf



add_docu = Document(r'%s\%s_report.docx' % (g_report.b.path_name, g_report.b.b_n))

se = add_docu.add_section()
se.left_margin = Inches(0.4)
se.right_margin = Inches(0.4)

# 添加第七节
# 如何操作
t8 = add_docu.add_heading('', 0).add_run(g_report.contents_lsit[6])
t8.font.size = Pt(25)  # 字体大小
t8.bold = True  # 字体是否加粗
t8.font.name = 'Times New Roman'  # 控制是西文时的字体
t8.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
p = add_docu.add_paragraph('')
p.paragraph_format.alignment = s.WD_ALIGN_PARAGRAPH.CENTER
t8.font.color.rgb = RGBColor(*g_report.color)

step1 = add_docu.add_paragraph('')
s1 = step1.add_run('1、扫描下方二维码或者微信搜索')
s1.font.size = Pt(17)  # 字体大小
s1.font.name = 'Times New Roman'  # 控制是西文时的字体
s1.element.rPr.rFonts.set(qn('w:eastAsia'), '楷体')

s2 = step1.add_run('”Pythonfan“')
s2.font.size = Pt(17)  # 字体大小
s2.bold = True  # 字体是否加粗
s2.font.name = 'Times New Roman'  # 控制是西文时的字体
s2.font.color.rgb = RGBColor(*g_report.color)

s3 = step1.add_run('，关注订阅号；')
s3.font.size = Pt(17)  # 字体大小
s3.font.name = 'Times New Roman'  # 控制是西文时的字体
s3.element.rPr.rFonts.set(qn('w:eastAsia'), '楷体')

tab = add_docu.add_table(rows=1, cols=3)  # 添加一个1行3列的空表
cell = tab.cell(0, 1)  # 获取某单元格对象（从0开始索引）
ph = cell.paragraphs[0]
run = ph.add_run()
run.add_picture(r'%s\anfany.jpg' % g_report.fig_path, width=Inches(2))

step1 = add_docu.add_paragraph('')
s1 = step1.add_run('2、回复关键词：')
s1.font.size = Pt(17)  # 字体大小
s1.font.name = 'Times New Roman'  # 控制是西文时的字体
s1.element.rPr.rFonts.set(qn('w:eastAsia'), '楷体')

s2 = step1.add_run('baby')
s2.font.size = Pt(17)  # 字体大小
s2.bold = True  # 字体是否加粗
s2.font.name = 'Times New Roman'  # 控制是西文时的字体
s2.font.color.rgb = RGBColor(*g_report.color)

s3 = step1.add_run('，获取微信号，添加好友，并备注：')
s3.font.size = Pt(17)  # 字体大小
s3.font.name = 'Times New Roman'  # 控制是西文时的字体
s3.element.rPr.rFonts.set(qn('w:eastAsia'), '楷体')

s2 = step1.add_run('报告')
s2.font.size = Pt(17)  # 字体大小
s2.bold = True  # 字体是否加粗
s2.font.name = 'Times New Roman'  # 控制是西文时的字体
s2.font.color.rgb = RGBColor(*g_report.color)

s3 = step1.add_run('；')
s3.font.size = Pt(17)  # 字体大小
s3.font.name = 'Times New Roman'  # 控制是西文时的字体
s3.element.rPr.rFonts.set(qn('w:eastAsia'), '楷体')

step1 = add_docu.add_paragraph('')
s1 = step1.add_run('3、按下面格式发送消息：')
s1.font.size = Pt(17)  # 字体大小
s1.font.name = 'Times New Roman'  # 控制是西文时的字体
s1.element.rPr.rFonts.set(qn('w:eastAsia'), '楷体')

step1 = add_docu.add_paragraph('', style='Intense Quote')
s1 = step1.add_run('宝宝昵称：（例，Eric）')
s1.font.size = Pt(15)  # 字体大小
s1.font.name = 'Times New Roman'  # 控制是西文时的字体
s1.element.rPr.rFonts.set(qn('w:eastAsia'), '华文中宋')
s1.font.color.rgb = RGBColor(*g_report.color)

step1 = add_docu.add_paragraph('', style='Intense Quote')
s1 = step1.add_run('性别：（例，男）')
s1.font.size = Pt(15)  # 字体大小
s1.font.name = 'Times New Roman'  # 控制是西文时的字体
s1.element.rPr.rFonts.set(qn('w:eastAsia'), '华文中宋')
s1.font.color.rgb = RGBColor(*g_report.color)

step1 = add_docu.add_paragraph('', style='Intense Quote')
s2 = step1.add_run('出生日期：（例，2016/11/22）')
s2.font.size = Pt(15)  # 字体大小
s2.font.name = 'Times New Roman'  # 控制是西文时的字体
s2.element.rPr.rFonts.set(qn('w:eastAsia'), '华文中宋')
s2.font.color.rgb = RGBColor(*g_report.color)

step1 = add_docu.add_paragraph('', style='Intense Quote')
s2 = step1.add_run('监测日期：（例：2016/11/22，2017/6/23，2018/10/10）')
s2.font.size = Pt(15)  # 字体大小
s2.font.name = 'Times New Roman'  # 控制是西文时的字体
s2.element.rPr.rFonts.set(qn('w:eastAsia'), '华文中宋')
s2.font.color.rgb = RGBColor(*g_report.color)

step1 = add_docu.add_paragraph('', style='Intense Quote')
s2 = step1.add_run('对应监测日期的身高（厘米）：（例：50.2，62.9，87.3）')
s2.font.size = Pt(15)  # 字体大小
s2.font.name = 'Times New Roman'  # 控制是西文时的字体
s2.element.rPr.rFonts.set(qn('w:eastAsia'), '华文中宋')
s2.font.color.rgb = RGBColor(*g_report.color)

step1 = add_docu.add_paragraph('', style='Intense Quote')
s2 = step1.add_run('对应监测日期的体重（千克）：（例：3.2，8.8，13.3）')
s2.font.size = Pt(15)  # 字体大小
s2.font.name = 'Times New Roman'  # 控制是西文时的字体
s2.element.rPr.rFonts.set(qn('w:eastAsia'), '华文中宋')
s2.font.color.rgb = RGBColor(*g_report.color)

step1 = add_docu.add_paragraph('', style='Intense Quote')
s2 = step1.add_run('对应监测日期的头围（厘米）：（例：34.6，39.5，49.4）')
s2.font.size = Pt(15)  # 字体大小
s2.font.name = 'Times New Roman'  # 控制是西文时的字体
s2.element.rPr.rFonts.set(qn('w:eastAsia'), '华文中宋')
s2.font.color.rgb = RGBColor(*g_report.color)

step1 = add_docu.add_paragraph('')
s1 = step1.add_run('4、发送消息后的72小时内可收到PDF报告。')
s1.font.size = Pt(17)  # 字体大小
s1.font.name = 'Times New Roman'  # 控制是西文时的字体
s1.element.rPr.rFonts.set(qn('w:eastAsia'), '楷体')

add_docu.add_paragraph('')
add_docu.add_paragraph('')
add_docu.add_paragraph('')
tab = add_docu.add_table(rows=1, cols=3)  # 添加一个1行3列的空表
cell = tab.cell(0, 1)  # 获取某单元格对象（从0开始索引）
ph = cell.paragraphs[0]
run = ph.add_run()
run.add_picture(r'%s\baby.jpg' % g_report.fig_path, width=Inches(4.6))

add_docu.add_paragraph('')
add_docu.add_paragraph('')
p = add_docu.add_paragraph('')
text1 = p.add_run('为了孩子，我的举动必须非常温和而慎重。')
text1.font.size = Pt(25)  # 字体大小
text1.bold = True  # 字体是否加粗
text1.font.name = 'Times New Roman'  # 控制是西文时的字体
text1.element.rPr.rFonts.set(qn('w:eastAsia'), '方正姚体')
p.paragraph_format.alignment = s.WD_ALIGN_PARAGRAPH.CENTER
text1.font.color.rgb = RGBColor(*g_report.color)

c = add_docu.add_paragraph('')
text1 = c.add_run('———卡尔.马克思')
text1.font.size = Pt(24)  # 字体大小
text1.bold = True  # 字体是否加粗
text1.font.name = 'Times New Roman'  # 控制是西文时的字体
text1.element.rPr.rFonts.set(qn('w:eastAsia'), '华文隶书')
c.paragraph_format.alignment = s.WD_ALIGN_PARAGRAPH.RIGHT
text1.font.color.rgb = RGBColor(*g_report.color)

add_docu.save(r'%s\%s_report.docx' % (g_report.b.path_name, g_report.b.b_n))


def doc2pdf(doc_name, pdf_name):
    """
    :word文件转pdf
    :param doc_name word文件名称
    :param pdf_name 转换后pdf文件名称
    """
    word = client.DispatchEx("Word.Application")
    worddoc = word.Documents.Open(doc_name, ReadOnly=1)
    worddoc.SaveAs(pdf_name, FileFormat=17)
    worddoc.Close()
    return pdf_name


doc2pdf(r'%s\%s_report.docx' % (g_report.b.path_name, g_report.b.b_n),
        r'%s\%s_report.pdf' % (g_report.b.path_name, g_report.b.b_n))


