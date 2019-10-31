# -*- coding：utf-8 -*-
# &Author  AnFany


# 根据宝宝信息，输出0-5岁宝宝身高、体重、头围、体重指数、身长别体重、身高别体重发育曲线监测报告

import stand_data_spyder as get_data # 获取数据
import baby_grow as plot_data
import os
from docx import Document
from docx.shared import Inches, RGBColor
from docx.oxml.ns import qn
from docx.shared import Pt
import docx.enum.text as s
import datetime
import time

fig_path = r'C:\Users\GWT9\Desktop\baby_data'

#  宝宝的信息
b = plot_data.Baby('Alina', '2015/11/22', 'girls',
                   ['2015/11/22', '2016/2/1', '2016/6/23', '2016/10/10', '2017/5/8', '2017/10/16',
                    '2018/4/3', '2018/9/9', '2019/3/18', '2019/9/11', '2020/4/11', '2020/8/26'],
                   [47.3, 50.2, 68.9, 74.3,  82.6, 90.6,
                    95.8, 96.3, 98.6, 100.7, 101.8, 108.6],
                   [2.2, 4.3,  6.9, 8.3, 10.6, 11.6,
                    12.3, 12.9, 13.6, 14.8, 15.8, 17.3],
                   [30.3, 33.7, 40.4, 43.4,  45.5, 49.1,
                    50.8, 51.3, 51.8, 52.3, 52.6, 52.9])

# 判断上面的文件夹中
c = b.b_n
if not os.path.exists(r'%s\%s' % (b.path_name, b.b_n)):
    # 新建一个文件夹子
    os.mkdir(r'%s\%s' % (b.path_name, c))
else:
    # 如果有重复，则区分开
    sign = 0
    while os.path.exists(r'%s\%s' % (b.path_name, c)):
        c = c[:len(b.b_n)] + str(sign)
        sign += 1
    os.mkdir(r'%s\%s' % (b.path_name, c))


# 更改后的文件夹
b.path_name += '/%s' % c

# 新建Word文档
document = Document()

# 颜色
if b.g == 'boys':
    color = (46, 92, 155)
else:
    color = (181, 62, 145)

#  封面制作
document.add_picture(r'%s\logo.png' % fig_path, width=Inches(2.5))

document.add_paragraph('')
p = document.add_paragraph('')
text1 = p.add_run('0-5岁%s童身体生长情况报告' % ('女' if b.g == 'girls' else '男'))
text1.font.size = Pt(24)                                # 字体大小
text1.bold = True                                       # 字体是否加粗
text1.font.name = 'Times New Roman'                     # 控制是西文时的字体
text1.element.rPr.rFonts.set(qn('w:eastAsia'), '华文中宋')
p.paragraph_format.alignment = s.WD_ALIGN_PARAGRAPH.CENTER
text1.font.color.rgb = RGBColor(*color)

p = document.add_paragraph('')
text2 = p.add_run('基于世界卫生组织发布的0-5岁%s童生长百分位数标准曲线制作' % ('女' if b.g == 'girls' else '男'))
text2.font.size = Pt(15)
text2.bold = False  # 字体是否加粗
text2.font.name = 'Times New Roman'
text2.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
p.paragraph_format.alignment = s.WD_ALIGN_PARAGRAPH.CENTER
text2.font.color.rgb = RGBColor(*color)

document.add_paragraph('')
tab = document.add_table(rows=1, cols=3)  # 添加一个1行3列的空表
cell = tab.cell(0, 1)  # 获取某单元格对象（从0开始索引）
ph = cell.paragraphs[0]
run = ph.add_run()
run.add_picture(r'%s\face.jpg' % fig_path, width=Inches(5))

document.add_paragraph('')
p = document.add_paragraph('')
key_id = b.generate_id()  # 生成报告的唯一编号，可回溯
text1 = p.add_run('报告编号：%s' % key_id)
text1.font.size = Pt(13)                                # 字体大小
text1.bold = True                                       # 字体是否加粗
text1.font.name = 'Times New Roman'                     # 控制是西文时的字体
text1.element.rPr.rFonts.set(qn('w:eastAsia'), '幼圆')
p.paragraph_format.alignment = s.WD_ALIGN_PARAGRAPH.CENTER
text1.font.color.rgb = RGBColor(*color)

current_date = datetime.datetime.now()
year = str(current_date.year)
month = str(current_date.month)
day = str(current_date.day)
p = document.add_paragraph('')
text1 = p.add_run('报告日期：%s年%s月%s日' % (year, (month if len(month) == 2 else '0' + month),
                                      (day if len(day) == 2 else '0' + day)))
text1.font.size = Pt(13)                                # 字体大小
text1.bold = True                                       # 字体是否加粗
text1.font.name = 'Times New Roman'                     # 控制是西文时的字体
text1.element.rPr.rFonts.set(qn('w:eastAsia'), '幼圆')
p.paragraph_format.alignment = s.WD_ALIGN_PARAGRAPH.CENTER
text1.font.color.rgb = RGBColor(*color)

# 目录页制作
t = document.add_heading('', 0).add_run('目录')
t.font.size = Pt(25)                                # 字体大小
t.bold = True                                       # 字体是否加粗
t.font.name = 'Times New Roman'                     # 控制是西文时的字体
t.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
p.paragraph_format.alignment = s.WD_ALIGN_PARAGRAPH.CENTER
t.font.color.rgb = RGBColor(*color)
document.add_paragraph('')
document.add_paragraph('')

if b.j_t == '0':
    two = '身长发育情况'
    six = '身长别体重情况'
    title_str = '身长'
elif b.j_t == '1':
    two = '身高发育情况'
    six = '身高别体重情况'
    title_str = '身高'
else:
    two = '身长、身高发育情况'
    six = '身长、身高别体重情况'
    title_str = '身长\高'

contents_lsit = ['一、宝宝信息表', '二、' + two, '三、体重发育情况',
                 '四、头围发育情况', '五、体重指数情况', '六、' + six, '七、生成宝宝专属报告']

for kk in contents_lsit:
    v = document.add_paragraph('', style='List Bullet')
    text1 = v.add_run(kk)
    text1.font.size = Pt(18)                                # 字体大小
    text1.font.name = 'Times New Roman'                     # 控制是西文时的字体
    text1.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    text1.bold = True
    text1.font.color.rgb = RGBColor(*color)
    document.add_paragraph('')

document.add_paragraph('')
document.add_paragraph('')
document.add_paragraph('')
document.add_paragraph('')

# 第一页：宝宝信息表
t = document.add_heading('', 0).add_run(contents_lsit[0])
t.font.size = Pt(25)                                # 字体大小
t.bold = True                                       # 字体是否加粗
t.font.name = 'Times New Roman'                     # 控制是西文时的字体
t.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
p.paragraph_format.alignment = s.WD_ALIGN_PARAGRAPH.CENTER
t.font.color.rgb = RGBColor(*color)

t = document.add_heading('', level=1).add_run('1、基础信息')
t.font.size = Pt(22)                                # 字体大小
t.bold = True                                       # 字体是否加粗
t.font.name = 'Times New Roman'                     # 控制是西文时的字体
t.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
t.font.color.rgb = RGBColor(*color)

if b.g == 'girls':
    style = 'Light List Accent 2'
else:
    style = 'Light List Accent 1'

table = document.add_table(rows=5, cols=2, style=style)

hdr_cells = table.rows[1].cells
hdr_cells[0].text = '宝宝昵称'
hdr_cells[1].text = b.b_n

hdr_cells = table.rows[2].cells
hdr_cells[0].text = '性别'
hdr_cells[1].text = '%s童' % ('女' if b.g == 'girls' else '男')

hdr_cells = table.rows[3].cells
hdr_cells[0].text = '出生日期'
y, m, d = str(b.b_d.year), str(b.b_d.month), str(b.b_d.day)
hdr_cells[1].text = '%s年%s月%s日' % (y, (m if len(m) == 2 else '0' + m),
                                      (d if len(d) == 2 else '0' + d))

keywords = '%s%s%s' % (y, (m if len(m) == 2 else '0' + m), (d if len(d) == 2 else '0' + d))

hdr_cells = table.rows[4].cells
hdr_cells[0].text = '年龄'
age_str = b.get_date_sub(b.b_d_str, time.strftime("%Y/%m/%d", time.localtime())).replace('$', '')
hdr_cells[1].text = age_str

for row in table.rows:
    for cell in row.cells:
        paragraphs = cell.paragraphs
        for paragraph in paragraphs:
            for run in paragraph.runs:
                font = run.font
                font.size = Pt(15)
                font.color.rgb = RGBColor(*color)


t2 = document.add_heading('', level=1).add_run('2、身体生长监测信息')
t2.font.size = Pt(22)                                # 字体大小
t2.bold = True                                       # 字体是否加粗
t2.font.name = 'Times New Roman'                     # 控制是西文时的字体
t2.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
t2.font.color.rgb = RGBColor(*color)

title_list = ['监测日期', title_str + '(厘米)', '体重(千克)', '头围(厘米)']


number_list = [b.o_d, b.o_l, b.o_w, b.o_hc]


if b.g == 'girls':
    style2 = 'Medium Shading 1 Accent 2'
else:
    style2 = 'Medium Shading 1 Accent 1'

table2 = document.add_table(rows=len(b.o_d)+1, cols=len(title_list), style=style)

for row in range(len(b.o_d) + 1):
    hdr_cells2 = table2.rows[row].cells
    if row == 0:
        for col in range(len(title_list)):
            hdr_cells2[col].text = title_list[col]
    else:
        for col in range(len(title_list)):
            try:
                hdr_cells2[col].text = str(number_list[col][row - 1])
            except IndexError:
                hdr_cells2[col].text = ''
# 改变网格的字体
sign = 1
for row2 in table2.rows:
    for cell in row2.cells:
        paragraphs = cell.paragraphs
        for paragraph in paragraphs:
            for run in paragraph.runs:
                font = run.font
                font.size = Pt(15)
                if not sign:
                    font.color.rgb = RGBColor(*color)
    sign = 0

# 加一个分页
document.add_page_break()

# WHO标准曲线的数据
D = get_data.DATA()
data_dict = D.read_data(b.g)

# 针对不同的项目绘制百分位线，按照顺序生成图片
date_list_dat = b.get_observe_days(b.o_d)
birth_two_date_str = []  # 0-2岁的字符串日期
birth_two_len_h = []  # 0-2岁的身长
birth_two_weight = []  # 0-2岁的体重
two_five_date_str = []  # 2-5岁的字符串日期
two_five_len_h = []  # 2-5岁的身高
two_five_weight = []  # 2-5岁的体重

len_h_doc = []
len_doc_b_t = []
h_doc_t_f = []
weight_doc = []

se = document.add_section()
se.left_margin = Inches(0)
se.right_margin = Inches(0)


for item in ['lhfa', 'wfa', 'hcfa', 'bfa', 'wfl', 'wfh']:
    if item == 'lhfa':  # 身长、身高
        # 需要区分年龄段
        for str_d, day_d, len_hei, wei in zip(b.o_d, date_list_dat, b.o_l, b.o_w):
            if day_d <= 730:  # 0-2岁的
                birth_two_date_str.append(str_d)
                birth_two_len_h.append(len_hei)
                birth_two_weight.append(wei)
            else:  # 2-5岁的
                two_five_date_str.append(str_d)
                two_five_len_h.append(len_hei)
                two_five_weight.append(wei)
        # 进行绘图
        if birth_two_date_str:
            len_doc_b_t = b.plot_p_data_lhfa(birth_two_len_h, birth_two_date_str,
                                             data_dict['p_exp'][item], '%s_%s0_2' % (b.g, item),
                                             start_y=0, end_y=2, min_num=0, max_num=730,
                                             fig_size=(6.7, 7), count_y=8, y_tap=(35, 105),
                                             data_loc=(0.45, 0.103, 0.35, 0.323),
                                             table_w=(0.125, 0.35, 0.35, 0.42, 0.32), max_table=12)
        if two_five_date_str:
            len_h25 = two_five_len_h.copy()
            h_doc_t_f = b.plot_p_data_lhfa(len_h25, two_five_date_str,
                                           data_dict['p_exp'][item], '%s_%s2_5' % (b.g, item),
                                           start_y=2, end_y=5, min_num=731, max_num=1856,
                                           fig_size=(6.7, 7), count_y=8, y_tap=(65, 135),
                                           data_loc=(0.4519, 0.103, 0.35, 0.318),
                                           table_w=(0.125, 0.35, 0.35, 0.42, 0.32), max_table=12)
        if h_doc_t_f:
            len_h_doc = len_doc_b_t + h_doc_t_f[1:]
        else:
            len_h_doc = len_doc_b_t

        # 身长、身高页
        t3 = document.add_heading('', 0).add_run(contents_lsit[1])
        t3.font.size = Pt(25)  # 字体大小
        t3.bold = True  # 字体是否加粗
        t3.font.name = 'Times New Roman'  # 控制是西文时的字体
        t3.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
        p.paragraph_format.alignment = s.WD_ALIGN_PARAGRAPH.CENTER
        t3.font.color.rgb = RGBColor(*color)
        if b.j_t == 0:
            # 添加0-2岁的
            tab = document.add_table(rows=1, cols=3)  # 添加一个1行3列的空表
            cell = tab.cell(0, 1)  # 获取某单元格对象（从0开始索引）
            ph = cell.paragraphs[0]
            run = ph.add_run()
            run.add_picture(r'%s\%s_%s0_2.png' % (b.path_name, b.g, item), width=Inches(7.6))
        elif b.j_t == 1:
            # 添加2-5岁的
            tab = document.add_table(rows=1, cols=3)  # 添加一个1行3列的空表
            cell = tab.cell(0, 1)  # 获取某单元格对象（从0开始索引）
            ph = cell.paragraphs[0]
            run = ph.add_run()
            run.add_picture(r'%s\%s_%s2_5.png' % (b.path_name, b.g, item), width=Inches(7.6))
        else:
            tab = document.add_table(rows=1, cols=3)  # 添加一个1行3列的空表
            cell = tab.cell(0, 1)  # 获取某单元格对象（从0开始索引）
            ph = cell.paragraphs[0]
            run = ph.add_run()
            run.add_picture(r'%s\%s_%s0_2.png' % (b.path_name, b.g, item), width=Inches(7.6))
            document.add_paragraph('')
            document.add_paragraph('')
            tab = document.add_table(rows=1, cols=3)  # 添加一个1行3列的空表
            cell = tab.cell(0, 1)  # 获取某单元格对象（从0开始索引）
            ph = cell.paragraphs[0]
            run = ph.add_run()
            run.add_picture(r'%s\%s_%s2_5.png' % (b.path_name, b.g, item), width=Inches(7.6))

    elif item == 'wfa':  # 体重
        weight_doc = b.plot_p_data_wfa(b.o_w, b.o_d, data_dict['p_exp'][item], '%s_%s0_5' % (b.g, item),
                                       start_y=0, end_y=5, min_num=0, max_num=1856,
                                       fig_size=(6.7, 7), count_y=11, y_tap=(1, 31),
                                       data_loc=(0.2138, 0.558, 0.35, 0.33),
                                       table_w=(0.125, 0.35, 0.335, 0.39, 0.31))
        # 体重页
        document.add_page_break()
        t4 = document.add_heading('', 0).add_run(contents_lsit[2])
        t4.font.size = Pt(25)  # 字体大小
        t4.bold = True  # 字体是否加粗
        t4.font.name = 'Times New Roman'  # 控制是西文时的字体
        t4.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
        p.paragraph_format.alignment = s.WD_ALIGN_PARAGRAPH.CENTER
        t4.font.color.rgb = RGBColor(*color)

        document.add_paragraph('')
        tab = document.add_table(rows=1, cols=3)  # 添加一个1行3列的空表
        cell = tab.cell(0, 1)  # 获取某单元格对象（从0开始索引）
        ph = cell.paragraphs[0]
        run = ph.add_run()
        run.add_picture(r'%s\%s_%s0_5.png' % (b.path_name, b.g, item), width=Inches(7.6))

    elif item == 'hcfa':  # 头围
        b.plot_p_data_hcfa(b.o_hc, b.o_d, data_dict['p_exp'][item], '%s_%s0_5' % (b.g, item),
                           start_y=0, end_y=5, min_num=0, max_num=1856,
                           fig_size=(6.7, 7), count_y=8, y_tap=(25, 60),
                           data_loc=(0.45, 0.103, 0.3523, 0.35), table_w=(0.125, 0.35, 0.325, 0.42, 0.33),
                           max_table=12)
        t5 = document.add_heading('', 0).add_run(contents_lsit[3])
        t5.font.size = Pt(25)  # 字体大小
        t5.bold = True  # 字体是否加粗
        t5.font.name = 'Times New Roman'  # 控制是西文时的字体
        t5.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
        p.paragraph_format.alignment = s.WD_ALIGN_PARAGRAPH.CENTER
        t5.font.color.rgb = RGBColor(*color)

        tab = document.add_table(rows=1, cols=3)  # 添加一个1行3列的空表
        cell = tab.cell(0, 1)  # 获取某单元格对象（从0开始索引）
        ph = cell.paragraphs[0]
        run = ph.add_run()
        run.add_picture(r'%s\%s_%s0_5.png' % (b.path_name, b.g, item), width=Inches(7.6))

    elif item == 'bfa':  # 体重指数
        if b.o_w[0]:
            length_list = b.o_w
        else:
            length_list = b.o_w[1:]
        bmi_list = [round((i / ((j / 100) ** 2)), 1) for i, j in zip(length_list, b.o_l)]
        b.plot_p_data_bfa(bmi_list, b.o_d, data_dict['p_exp'][item], len_h_doc, weight_doc,
                          '%s_%s0_5' % (b.g, item),
                          start_y=0, end_y=5, min_num=0, max_num=1856,
                          fig_size=(6.7, 7), count_y=12, y_tap=(4, 26),
                          data_loc=(0.3598, 0.103, 0.35, 0.33),
                          table_w=(0.125, 0.33, 0.325, 0.41, 0.41, 0.41, 0.26),
                          max_table=12)

        t6 = document.add_heading('', 0).add_run(contents_lsit[4])
        t6.font.size = Pt(25)  # 字体大小
        t6.bold = True  # 字体是否加粗
        t6.font.name = 'Times New Roman'  # 控制是西文时的字体
        t6.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
        p.paragraph_format.alignment = s.WD_ALIGN_PARAGRAPH.CENTER
        t6.font.color.rgb = RGBColor(*color)

        tab = document.add_table(rows=1, cols=3)  # 添加一个1行3列的空表
        cell = tab.cell(0, 1)  # 获取某单元格对象（从0开始索引）
        ph = cell.paragraphs[0]
        run = ph.add_run()
        run.add_picture(r'%s\%s_%s0_5.png' % (b.path_name, b.g, item), width=Inches(7.6))

        # 身长、身高别体重页
        t7 = document.add_heading('', 0).add_run(contents_lsit[5])
        t7.font.size = Pt(25)  # 字体大小
        t7.bold = True  # 字体是否加粗
        t7.font.name = 'Times New Roman'  # 控制是西文时的字体
        t7.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
        p.paragraph_format.alignment = s.WD_ALIGN_PARAGRAPH.CENTER
        t7.font.color.rgb = RGBColor(*color)

    elif item == 'wfl':  # 身长别体重
        length_list = []
        if birth_two_len_h:
            if birth_two_len_h[0]:
                length_list = birth_two_len_h
            else:
                length_list = birth_two_len_h[1:]

            b.plot_p_data_wfl(birth_two_weight, length_list, birth_two_date_str, data_dict['p_exp'][item],
                              len_doc_b_t, weight_doc[:len(len_doc_b_t)],
                              '%s_%s0_2' % (b.g, item), start_y=0, end_y=2, min_num=45, max_num=110,
                              fig_size=(6.7, 7), count_y=17, y_tap=(1, 33),
                              data_loc=(0.355, 0.539, 0.36, 0.35),
                              table_w=(0.125, 0.32, 0.37, 0.41, 0.41, 0.39, 0.26))
            # 添加0-2岁的
            tab = document.add_table(rows=1, cols=3)  # 添加一个1行3列的空表
            cell = tab.cell(0, 1)  # 获取某单元格对象（从0开始索引）
            ph = cell.paragraphs[0]
            run = ph.add_run()
            run.add_picture(r'%s\%s_%s0_2.png' % (b.path_name, b.g, item), width=Inches(7.6))

    else:  # 身高别体重
        if two_five_len_h:
            b.plot_p_data_wfh(two_five_weight, two_five_len_h, two_five_date_str, data_dict['p_exp'][item],
                              h_doc_t_f, [weight_doc[0]] + weight_doc[len(len_doc_b_t):],
                              '%s_%s2_5' % (b.g, item), start_y=2, end_y=5, min_num=65, max_num=120,
                              fig_size=(6.7, 7), count_y=17, y_tap=(4, 36),
                              data_loc=(0.355, 0.539, 0.36, 0.35),
                              table_w=(0.125, 0.32, 0.37, 0.41, 0.41, 0.39, 0.26))
            # 添加2-5岁的
            document.add_paragraph('')
            document.add_paragraph('')
            tab = document.add_table(rows=1, cols=3)  # 添加一个1行3列的空表
            cell = tab.cell(0, 1)  # 获取某单元格对象（从0开始索引）
            ph = cell.paragraphs[0]
            run = ph.add_run()
            run.add_picture(r'%s\%s_%s2_5.png' % (b.path_name, b.g, item), width=Inches(7.6))
            document.add_page_break()



document.save(r'%s\%s_report.docx' % (b.path_name, b.b_n))